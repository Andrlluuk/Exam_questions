import math
import os
import random
import shutil
import subprocess

from PIL import Image
from docx import Document
from fpdf import FPDF
from pdf2image import convert_from_path

from .parser import parse_tex, parse_docx, parse_txt
from .statistics import collect_statistics

from pylatex import Document
from .models import File


def get_minimal_number_of_questions(questions_pool, params, stats):
    number_of_questions = {}
    for key in questions_pool.keys():
        for mark in questions_pool[key].keys():
            number_of_questions[mark] = 0
    for key in questions_pool.keys():
        for mark in questions_pool[key].keys():
            if '0' in questions_pool[key][mark].keys():
                number_of_questions[mark] += len(questions_pool[key][mark]['0'])

    for key in number_of_questions.keys():
        if params[key] - stats['density'][key] == 0:
            continue
        number_of_questions[key] /= (params[key] - stats['density'][key])

    return math.ceil(max(number_of_questions.values()))


def chunk(L, n=1, verbose=False):
    total = len(L)
    if n > 0:
        size = int(total / n)
        rest = total % n
        if verbose:
            msg = "{} items to be split into {} chunks of size {} with {} extra"
            print
            msg.format(total, n, size, rest)
        if not size:
            return [[x] for x in L] + [[] for i in range(n - total)]
        if rest:
            index = [x for x in range(0, total, size)]
            extra = [index[i] + i for i in range(rest + 1)] + \
                    [x + rest for x in index[rest + 1:][:n - rest]]
            ranges = [(extra[i], extra[i + 1]) for i in range(len(extra) - 1)]
        else:
            index = [x for x in range(0, total + 1, size)]
            ranges = [(index[i], index[i + 1]) for i in range(len(index) - 1)]
        return [L[i:j] for i, j in ranges]


def indexes_ok(indexes, params, info_tickets, mark):
    for idx in indexes:
        if (info_tickets[idx][mark] + 1 > params[mark]):
            return False

    return True


def create_questions_tex(path, filename, params, questions_pool, tickets):
    """TODO: change"""
    info_tickets = {i + 1: {'3': 0, '4': 0, '5': 0, '6': 0} for i in range(tickets)}
    dict_of_tickets = {i + 1: [] for i in range(tickets)}
    for chapter in questions_pool.keys():
        for mark in questions_pool[chapter].keys():
            for necessarity in questions_pool[chapter][mark].keys():
                if necessarity == '0':
                    continue
                else:
                    for question in questions_pool[chapter][mark][necessarity]:
                        possible = set([i for i in range(tickets) if params[mark] - info_tickets[i + 1][mark] > 0])
                        indexes = random.sample(possible, math.floor(tickets / int(necessarity)))
                        for idx in indexes:
                            info_tickets[idx + 1][mark] += 1
                            if mark == '6':
                                dict_of_tickets[idx + 1].append('Задача' + question)
                            else:
                                dict_of_tickets[idx + 1].append(str(mark) + question)
    unnecessary_pool = {'3': [], '4': [], '5': [], '6': []}
    for chapter in questions_pool.keys():
        for mark in questions_pool[chapter].keys():
            if '0' in questions_pool[chapter][mark]:
                unnecessary_pool[mark].extend(questions_pool[chapter][mark]['0'])

    for mark in unnecessary_pool.keys():
        curr = 1
        full = False
        while not full:
            full = True
            random.shuffle(unnecessary_pool[mark])
            for question in unnecessary_pool[mark]:
                if info_tickets[curr][mark] < params[mark] and question not in dict_of_tickets[curr]:
                    full = False
                    info_tickets[curr][mark] += 1
                    if mark == '6':
                        dict_of_tickets[curr].append('Задача' + question)
                    else:
                        dict_of_tickets[curr].append(str(mark) + question)
                curr = (curr) % tickets + 1
    return dict_of_tickets


def get_statistics(current_file, additional_file=None):
    title = None
    if additional_file == None:
        _, file_extension = os.path.splitext(current_file)
        if file_extension == '.tex':
            questions_pool, title = parse_tex(current_file)
        elif file_extension == '.docx' or file_extension == '.docx':
            questions_pool, title = parse_docx(current_file)
        elif file_extension == '.txt':
            questions_pool, title = parse_txt(current_file)
        else:
            return "Error"
        stat = collect_statistics(questions_pool)
        return stat, questions_pool, title
    else:
        _, file_extension1 = os.path.splitext(current_file)
        _, file_extension2 = os.path.splitext(additional_file)
        if file_extension1 == '.tex':
            questions_pool, title = parse_tex(current_file)
        elif file_extension1 == '.docx' or file_extension1 == '.docx':
            questions_pool = parse_docx(current_file)
        elif file_extension1 == '.txt':
            questions_pool = parse_txt(current_file)
        else:
            return "Error"
        if file_extension1 == '.tex':
            additional_questions_pool, _ = parse_tex(additional_file)
        elif file_extension1 == '.docx' or file_extension1 == '.docx':
            additional_questions_pool = parse_docx(additional_file)
        elif file_extension1 == '.txt':
            additional_questions_pool = parse_txt(additional_file)
        else:
            return "Error"
        pool = {**questions_pool, **additional_questions_pool}
        stat = collect_statistics(pool)
        return stat, pool, title


def remove_newline_signs(line):
    while (line.endswith('\n')) or (line.endswith('\\')) or (line.endswith('  ')):
        line = line[:-2]
    return line


def parsess_doc(folder, name, params):
    dir_path = ""
    filename = os.path.join(dir_path, name)
    doc = Document(filename)
    all_questions = []
    questions_pool = {}
    part_number = 0
    questions_pool[part_number] = {3: [], 4: [], 5: [], 6: []}
    for line in doc.paragraphs:
        line = line.text
        if len(line) == 0:
            continue
        if line[0] == '%':
            continue
        if 'Глава' in line:
            part_number += 1
            questions_pool[part_number] = {3: [], 4: [], 5: [], 6: []}
        if line.startswith(params['label_3']):
            if not params['show']:
                if ".png" in params['label_3'] or ".jpg" in params['label_3'] or ".jpeg" in params['label_3']:
                    pos = line.find(params['label_3'])
                    line = line[pos + len(params['label_3']) + 2:]
                else:
                    line = line.replace(params['label_3'], "")
                for i, x in enumerate(line):
                    if x.isalpha():  # True if its a letter
                        po = i  # first letter position
                        break

                line = line[po:]
            questions_pool[part_number][3].append(line)
        elif line.startswith(params['label_4']):
            if not params['show']:
                if ".png" in params['label_4'] or ".jpg" in params['label_4'] or ".jpeg" in params['label_4']:
                    pos = line.find(params['label_4'])
                    line = line[pos + len(params['label_4']) + 2:]
                else:
                    line = line.replace(params['label_4'], "")
                for i, x in enumerate(line):
                    if x.isalpha():  # True if its a letter
                        po = i  # first letter position
                        break

                line = line[po:]
            questions_pool[part_number][4].append(line)
        elif line.startswith(params['label_5']):
            if not params['show']:
                if ".png" in params['label_5'] or ".jpg" in params['label_5'] or ".jpeg" in params['label_5']:
                    pos = line.find(params['label_5'])
                    line = line[pos + len(params['label_5']) + 2:]
                else:
                    line = line.replace(params['label_5'], "")
                for i, x in enumerate(line):
                    if x.isalpha():  # True if its a letter
                        po = i  # first letter position
                        break

                line = line[po:]
            questions_pool[part_number][5].append(line)
        elif line.startswith(params['label_problem']):
            questions_pool[part_number][6].append(line)

    return questions_pool, dir_path


def parsess_tex(folder, name, params):
    # os.chdir(folder)
    dir_path = os.path.dirname(folder)
    filename = os.path.join(dir_path, name)
    title = []
    with open(filename, encoding='UTF-8') as f:
        all_questions = f.readlines()
    idx = 0
    while 'begin{document}' not in all_questions[idx]:
        if ('documentclass' in all_questions[idx]) or ('documentarticle' in all_questions[idx]):
            idx += 1
            continue
        title.append(all_questions[idx])
        idx += 1

    po = 0
    # creation
    questions_pool = {}
    part_number = 0
    for line in all_questions:
        if len(line) == 0:
            continue
        if line[0] == '%':
            continue
        if 'Глава' in line:
            part_number += 1
            questions_pool[part_number] = {3: [], 4: [], 5: [], 6: []}
        if line.startswith(params['label_3']):
            if not params['show']:
                if ".png" in params['label_3'] or ".jpg" in params['label_3'] or ".jpeg" in params['label_3']:
                    pos = line.find(params['label_3'])
                    line = line[pos + len(params['label_3']) + 2:]
                else:
                    line = line.replace(params['label_3'], "")
                po = 0
                for i, x in enumerate(line):
                    if x.isalpha():  # True if its a letter
                        po = i  # first letter position
                        break

                line = line[po:]
            questions_pool[part_number][3].append('3' + line)
        elif line.startswith(params['label_4']):
            if not params['show']:
                if ".png" in params['label_4'] or ".jpg" in params['label_4'] or ".jpeg" in params['label_4']:
                    pos = line.find(params['label_4'])
                    line = line[pos + len(params['label_4']) + 2:]
                else:
                    line = line.replace(params['label_4'], "")
                po = 0
                for i, x in enumerate(line):
                    if x.isalpha():  # True if its a letter
                        po = i  # first letter position
                        break

                line = line[po:]
            questions_pool[part_number][4].append('4' + line)
        elif line.startswith(params['label_5']):
            if not params['show']:
                if ".png" in params['label_5'] or ".jpg" in params['label_5'] or ".jpeg" in params['label_5']:
                    pos = line.find(params['label_5'])
                    line = line[pos + len(params['label_5']) + 2:]
                else:
                    line = line.replace(params['label_5'], "")
                po = 0
                for i, x in enumerate(line):
                    if x.isalpha():  # True if its a letter
                        po = i  # first letter position
                        break

                line = line[po:]
            questions_pool[part_number][5].append('5' + line)
        elif line.startswith(params['label_problem']):
            questions_pool[part_number][6].append('Задача' + line)

    return questions_pool, title


def create_texs(questions_pool, params, dir_path, title=[]):
    for ticket in questions_pool.keys():
        questions = questions_pool[ticket]
        with open(os.path.join(dir_path, f"tickets{ticket}.tex"), 'w') as f:
            f.write('\\documentclass[preview]{standalone} \n')
            if title != []:
                for line in title:
                    f.write('%s\n' % line)
            else:
                f.write('\\usepackage[english, russian]{babel} \n')
            f.write('\\pageclass{empty} \n')
            f.write('\\begin{document} \n')
            f.write('\\begin{center} {\\Large Билет №%s} \\end{center} \n' % str(ticket))
            f.write('\n')
            for i in range(len(questions)):
                line = remove_newline_signs(questions[i])
                if not params['show']:
                    line = line.split(' ', 1)[1]
                    f.write('%s' % f'{i + 1}. ' + line)
                else:
                    f.write('%s' % line)
                f.write('\\\\\n')
                f.write('\n')
            f.write('\\end{document}')
        os.chdir(dir_path)
        #doc = Document(f"{dir_path}/tickets{ticket + 1}.tex")
        #doc.generate_pdf(f"{dir_path}/tickets{ticket + 1}")
        # os.system(f"sudo pdflatex tickets{ticket + 1}.tex")
 
        subprocess.call(['pdflatex', '-interaction=nonstopmode', f"tickets{ticket}.tex"], stdout=subprocess.DEVNULL)
        os.chdir('../../../..')


def create_pdf(tickets, dir_path, title=[]):
    for ticket in range(tickets):
        file = convert_from_path(os.path.join(dir_path, f'tickets{ticket + 1}.pdf'), 500)
        for fil in file:
            fil.save(os.path.join(dir_path, f'tickets{ticket + 1}.png'), 'PNG')

    pdf = FPDF()
    x_current = 20
    y_current = 5
    max_height = 0
    for image in [os.path.join(dir_path, f'tickets{ticket + 1}.png') for ticket in range(tickets)]:
        im = Image.open(image)
        width, height = im.size
        coef = width / 140
        max_height = max(height / coef, max_height)
    pdf.add_page()
    for image in [os.path.join(dir_path, f'tickets{ticket + 1}.png') for ticket in range(tickets)]:
        im = Image.open(image)
        width, height = im.size
        coef = width / 140
        if y_current + max_height > 292:
            y_current = 5
            pdf.add_page()
        y_current += 5
        pdf.image(image, x_current, y_current, width / coef, height / coef)
        y_current += max_height + 5
        pdf.line(0, y_current, 297, y_current)
    pdf.output(os.path.join(dir_path, "tickets.pdf"), "F")
    # os.chdir('../../..')


def build_questions_from_tex(folder, name, params):
    questions_pool, dir_path, title = parse_tex(folder, name, params)
    if (params['additional_file']):
        _, file_extension = os.path.splitext(params['additional_file'])
        if file_extension == ".tex":
            questions_pool_additional, _, _ = parse_tex(folder, params['additional_file'], params)
        elif file_extension == ".doc" or file_extension == ".docx":
            questions_pool_additional, _ = parse_docx(folder, params['additional_file'], params)
        for i in range(1, min(len(questions_pool), len(questions_pool_additional))):
            for key in questions_pool[i].keys():
                questions_pool[i][key].extend(questions_pool_additional[i][key])
    if create_texs(questions_pool, params, dir_path, folder, title) == "Error":
        return "Error"
    create_pdf(questions_pool, params, dir_path, folder, title)


def doc_parsing(folder, name, params):
    questions_pool, dir_path = parse_docx(folder, name, params)
    if (params['additional_file']):
        _, file_extension = os.path.splitext(params['additional_file'])
        if file_extension == ".tex":
            questions_pool_additional, _, _ = parse_tex(folder, params['additional_file'], params)
        elif file_extension == ".doc" or file_extension == ".docx":
            questions_pool_additional, _ = parse_docx(folder, params['additional_file'], params)
        for i in range(1, min(len(questions_pool), len(questions_pool_additional))):
            for key in questions_pool[i].keys():
                questions_pool[i][key].extend(questions_pool_additional[i][key])
    if create_texs(questions_pool, params, dir_path, folder) == "Error":
        return "Error"
    create_pdf(questions_pool, params, dir_path, folder)


def create_folder(tickets, dir_path, title=[]):
    os.mkdir(os.path.join(dir_path, "ticket_folder"))
    for ticket in range(tickets):
        file = convert_from_path(os.path.join(dir_path, f'tickets{ticket + 1}.pdf'), 500)
        for fil in file:
            fil.save(os.path.join(dir_path, f'tickets{ticket + 1}.png'), 'PNG')

    for ticket in range(tickets):
        pdf = FPDF()
        x_current = 20
        y_current = 5
        max_height = 0
        image = os.path.join(dir_path, f'tickets{ticket + 1}.png')
        im = Image.open(image)
        width, height = im.size
        coef = width / 140
        max_height = max(height / coef, max_height)
        pdf.add_page()
        pdf.image(image, x_current, y_current, width / coef, height / coef)
        y_current += max_height
        y_current += 5
        pdf.line(0, y_current, 297, y_current)
        pdf.output(os.path.join(dir_path, f"ticket_folder/ticket_{ticket + 1}.pdf"), "F")
    shutil.make_archive(os.path.join(dir_path, f"ticket_folder"), 'zip', os.path.join(dir_path, f"ticket_folder"))

